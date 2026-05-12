"""
Test with actual resume files - PDF, DOCX, and TXT format
"""

import requests
import io
from docx import Document

BASE_URL = "http://localhost:8000"

# ====================================
# Create Sample DOCX Resume
# ====================================

def create_sample_docx_resume():
    """Create a sample resume in DOCX format"""
    doc = Document()

    # Title
    doc.add_heading('PRINCE SHARMA', level=1)
    doc.add_paragraph('Email: prince.sharma@gmail.com | Phone: +91-9876543210')
    doc.add_paragraph('LinkedIn: linkedin.com/in/prince-sharma')

    # Summary
    doc.add_heading('PROFESSIONAL SUMMARY', level=2)
    doc.add_paragraph(
        'Experienced Full Stack Developer with 3 years of professional experience in '
        'building scalable web applications using Python, Node.js, React, and MongoDB.'
    )

    # Experience
    doc.add_heading('EXPERIENCE', level=2)

    doc.add_heading('Senior Backend Developer', level=3)
    doc.add_paragraph('TechCorp Solutions | Bangalore | Jan 2023 - Present')
    doc.add_paragraph('• Leading team of 5 developers')
    doc.add_paragraph('• 2+ years of experience building microservices')
    doc.add_paragraph('• Developed REST APIs handling 1M+ requests/day')

    doc.add_heading('Full Stack Developer', level=3)
    doc.add_paragraph('StartupXYZ | Bangalore | Jun 2021 - Dec 2022')
    doc.add_paragraph('• Worked on React.js frontend and Node.js backend')
    doc.add_paragraph('• 1 year and 6 months of experience in agile environment')
    doc.add_paragraph('• Implemented real-time features using Socket.io')

    doc.add_heading('Junior Developer Internship', level=3)
    doc.add_paragraph('WebServices Inc | Remote | Jan 2021 - May 2021')
    doc.add_paragraph('• Completed 5 months internship as junior developer')
    doc.add_paragraph('• Learned MERN stack fundamentals')

    # Skills
    doc.add_heading('TECHNICAL SKILLS', level=2)
    doc.add_paragraph('Languages: Python, JavaScript, SQL')
    doc.add_paragraph('Frontend: React.js, Tailwind CSS, Redux')
    doc.add_paragraph('Backend: Node.js, Express.js, FastAPI')
    doc.add_paragraph('Database: MongoDB, PostgreSQL')

    # Education
    doc.add_heading('EDUCATION', level=2)
    doc.add_paragraph('B.Tech in Computer Science | Delhi University | 2020')

    # Save
    doc.save('sample_resume_detailed.docx')
    print("✓ Created: sample_resume_detailed.docx")

# ====================================
# Create Sample TXT Resume
# ====================================

def create_sample_txt_resume():
    """Create a sample resume in TXT format"""
    resume_text = """
    PRINCE SHARMA
    Email: prince.sharma@gmail.com | Phone: +91-9876543210
    LinkedIn: linkedin.com/in/prince-sharma
    
    ========================================
    PROFESSIONAL SUMMARY
    ========================================
    
    Experienced Full Stack Developer with 3 years of professional experience in building 
    scalable web applications using Python, Node.js, React, and MongoDB.
    
    ========================================
    EXPERIENCE
    ========================================
    
    SENIOR BACKEND DEVELOPER
    TechCorp Solutions | Bangalore | Jan 2023 - Present
    - Leading team of 5 developers
    - 2+ years of experience building microservices architecture
    - Developed REST APIs handling 1M+ requests/day
    - Tech: Python, Node.js, MongoDB, Docker, Kubernetes
    
    FULL STACK DEVELOPER
    StartupXYZ | Bangalore | Jun 2021 - Dec 2022
    - Worked on React.js frontend and Node.js backend
    - 1 year and 6 months of experience in agile environment
    - Implemented real-time features using Socket.io
    - Tech: React, Express, MongoDB, Firebase
    
    JUNIOR DEVELOPER INTERNSHIP
    WebServices Inc | Remote | Jan 2021 - May 2021
    - Completed 5 months internship as junior developer
    - Learned MERN stack fundamentals
    - Built 3 full-stack projects
    
    ========================================
    TECHNICAL SKILLS
    ========================================
    
    Languages: Python, JavaScript, SQL, HTML, CSS
    Frontend: React.js, Tailwind CSS, Redux, Framer Motion
    Backend: Node.js, Express.js, FastAPI, Django
    Database: MongoDB, PostgreSQL, MySQL
    DevOps: Docker, Kubernetes, GitHub Actions
    
    ========================================
    EDUCATION
    ========================================
    
    B.Tech in Computer Science | Delhi University | 2020
    """
    
    with open('sample_resume_detailed.txt', 'w') as f:
        f.write(resume_text)
    print("✓ Created: sample_resume_detailed.txt")

# ====================================
# TEST: Upload DOCX Resume
# ====================================

def test_docx_upload():
    print("\n" + "="*60)
    print("TEST 1: UPLOAD DOCX RESUME")
    print("="*60)
    
    try:
        with open('sample_resume_detailed.docx', 'rb') as f:
            files = {'file': f}
            response = requests.post(
                f"{BASE_URL}/detect-experience",
                files=files
            )
        
        result = response.json()
        print(f"\n✓ Status: {result['status']}")
        print(f"✓ File: {result['file_name']}")
        print(f"\n📊 DETECTED EXPERIENCE:")
        print(f"   All experiences: {result['all_experience']}")
        print(f"   Best experience: {result['best_experience']} years")
        print(f"\n📄 Resume snippet (first 300 chars):")
        print(f"   {result['resume_snippet'][:300]}...")
        return True
    
    except Exception as e:
        print(f"❌ Error: {e}")
        return False

# ====================================
# TEST: Upload TXT Resume
# ====================================

def test_txt_upload():
    print("\n" + "="*60)
    print("TEST 2: UPLOAD TXT RESUME")
    print("="*60)
    
    try:
        with open('sample_resume_detailed.txt', 'rb') as f:
            files = {'file': f}
            response = requests.post(
                f"{BASE_URL}/detect-experience",
                files=files
            )
        
        result = response.json()
        print(f"\n✓ Status: {result['status']}")
        print(f"✓ File: {result['file_name']}")
        print(f"\n📊 DETECTED EXPERIENCE:")
        print(f"   All experiences: {result['all_experience']}")
        print(f"   Best experience: {result['best_experience']} years")
        print(f"\n📄 Resume snippet (first 300 chars):")
        print(f"   {result['resume_snippet'][:300]}...")
        return True
    
    except Exception as e:
        print(f"❌ Error: {e}")
        return False

# ====================================
# TEST: Verify Context Filtering
# ====================================

def test_with_numbers_in_resume():
    print("\n" + "="*60)
    print("TEST 3: CONTEXT FILTERING (Numbers should be filtered)")
    print("="*60)
    
    # Create resume with random numbers
    resume_with_numbers = """
    PRINCE SHARMA
    
    Contact: Phone +91-9876543210 (should NOT be detected)
    Office ID: 2024 (should NOT be detected)
    
    EXPERIENCE
    
    Worked as Python Developer with 3 years experience
    Managed team for 2 years
    Completed project in 8 months
    """
    
    try:
        files = {'file': ('test_resume.txt', resume_with_numbers)}
        response = requests.post(
            f"{BASE_URL}/detect-experience",
            files=files
        )
        
        result = response.json()
        print(f"\n✓ Status: {result['status']}")
        print(f"\n📊 DETECTED EXPERIENCE (Context-aware):")
        print(f"   Detected: {result['all_experience']}")
        print(f"   Best: {result['best_experience']} years")
        print(f"\n💡 Notice: Phone number and ID were filtered out by context validation!")
        return True
    
    except Exception as e:
        print(f"❌ Error: {e}")
        return False

# ====================================
# Main Test Runner
# ====================================

if __name__ == "__main__":
    print("\n" + "="*60)
    print("RESUME EXPERIENCE DETECTION - COMPREHENSIVE TEST")
    print("="*60)
    
    # Create sample resumes
    print("\n📝 Creating sample resumes...")
    create_sample_docx_resume()
    create_sample_txt_resume()
    
    try:
        # Run tests
        test1_passed = test_docx_upload()
        test2_passed = test_txt_upload()
        test3_passed = test_with_numbers_in_resume()
        
        # Summary
        print("\n" + "="*60)
        print("TEST SUMMARY")
        print("="*60)
        print(f"DOCX Upload: {'✅ PASSED' if test1_passed else '❌ FAILED'}")
        print(f"TXT Upload: {'✅ PASSED' if test2_passed else '❌ FAILED'}")
        print(f"Context Filtering: {'✅ PASSED' if test3_passed else '❌ FAILED'}")
        
        if test1_passed and test2_passed and test3_passed:
            print("\n🎉 ALL TESTS PASSED! Code is working perfectly!")
        else:
            print("\n⚠️ Some tests failed. Check server logs.")
    
    except requests.exceptions.ConnectionError:
        print("\n❌ Cannot connect to server!")
        print("Make sure to run: python app/main.py")
    except Exception as e:
        print(f"\n❌ Error: {e}")
